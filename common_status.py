#!/usr/bin/python3
# -*- coding: utf-8 -*-

import sys
import os
import pickle
import logging
import datetime
import time
import stat

import docx
from docx.shared import Cm
from docx.shared import Pt

from PyQt5.QtWidgets import (QWidget, QTextEdit, QLineEdit, QPushButton, QScrollArea, QScrollBar, QSizePolicy,
    QHBoxLayout, QVBoxLayout, QApplication, QLabel, QRadioButton, QButtonGroup, QComboBox, QInputDialog, 
    QGridLayout, QGroupBox, QMessageBox, QDesktopWidget, QFileDialog, QTabWidget, QDateEdit, QListWidget, QTableWidget, QTableWidgetItem)
from PyQt5.QtCore import QVariant, QEvent, QPoint, QSize
from PyQt5.QtCore import (Qt, pyqtSignal, QRect, QVariantAnimation,QEasingCurve, QAbstractAnimation, 
                        QPropertyAnimation, pyqtProperty)
from PyQt5.QtGui import QFont, QFontMetrics, QPainter ,QPixmap, QImage, QColor, QPen, QCursor, QTextCharFormat, QTextCursor, QIcon
from common_notation import common_status_notation



class PushButton(QPushButton):
    

    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        
        #if self.color_ap == "":
        #    self._update_stylesheet(QColor(self.color_main), QColor("black"))
        #else:
        #    self._update_stylesheet(QColor(self.color_ap), QColor("black"))
        self.parent = parent
        self._color = "#ffb4a2"
        self._color_new = "#A64A35"
        self._font_size = 14
        self._child_windows = False
        self.setCursor(QCursor(Qt.PointingHandCursor))
        
        self._animation = QVariantAnimation(
            startValue=QColor(self._color_new),
            endValue=QColor(self._color),
            valueChanged=self._on_value_changed,
            duration=400,
        )

    @pyqtProperty(str)
    def color(self):
        return self._color

    @color.setter
    def color(self, value):
        self._color = value
        color_new = ""
        if value == "#f0f0f0": 
            color_new = "#B4B4B4"
        
        if value == "#B4B4B4": 
            color_new = "#D2D2D2"

        elif value == "#adff2f":
            color_new = "#6AA60F"
        elif value =="#ffb4a2":
            color_new = "#A64A35"

        elif value =="#A64A35":
            color_new = "#ffb4a2"

        elif value =="":
            color_new = ""
        elif value == "#fa8090":
            color_new = "#BB7780"
        
        elif value == "#FF7373":
            color_new = "#BF3030"

        elif value == "#FF4040":
            color_new = "#BF3030"

        elif value == "#BF3030":
            color_new = "#FF7373"


        self._color_new = color_new
        self._update_stylesheet(QColor(self._color),QColor("black"))

    @pyqtProperty(str)
    def color_new(self):
        return self._color_new

    @color_new.setter
    def color_new(self, color):
        if color == "#f0f0f0": 
            color_new = "#B4B4B4"
        elif color == "#adff2f":
            color_new = "#91BF4A"
        elif color =="#ffb4a2":
            color_new = "#A64A35"
        elif color =="":
            color_new = ""
        self._color_new = color_new

    @pyqtProperty(bool)
    def childWindows(self):
        return self._child_windows

    @childWindows.setter
    def childWindows(self, value):
        self._child_windows = value

    def setColor (self, color):
        
        return color
        #self._update_stylesheet(QColor(self.color_main),QColor("black"))
        #self.update()

    @pyqtProperty(int)
    def font_size(self):
        return self._font_size

    @font_size.setter
    def font_size(self, font_size):
        self._font_size = font_size

    def getColor (self):
        return self.color_main

    def _on_value_changed(self):
        foreground = (
            QColor("black")
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor("white")
        )
        background = (
            QColor(self._color)
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor(self._color_new)
            )
        self._update_stylesheet(background, foreground)


    def _update_stylesheet(self, background, foreground):

        self.setStyleSheet(
            """
        QPushButton{
            background-color: %s;
            border: none;
            color: %s;
            padding: 4px 8px;
            text-align: center;
            text-decoration: none;
            font-size: %spx;
            border-radius: 8px;
            margin: 4px 2px;
            min-height: 20 px;
            
        }
        """
            % (background.name(), foreground.name(), self._font_size)
        )

    def enterEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Backward)
        self._animation.start()
        super().enterEvent(event)

    def leaveEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Forward)
        self._animation.start()
        super().leaveEvent(event)

    def mousePressEvent(self, event):
        
        if self._child_windows:
            msg = QMessageBox()
            msg.setWindowTitle("Информация")
            msg.setText("Закончите работу с открытым окном!!!")
            msg.setIcon(QMessageBox.Information)
            center = [QDesktopWidget().screenGeometry(-1).width()/2, QDesktopWidget().screenGeometry(-1).height()/2] 
            msg.show()
            msg.move(center[0],center[1])
            msg.exec_()
            print('!!!!!!!!!mousePressEvent!!!!!!!!!!!!!!!!!!', event.x(), event.y(),)
            return None
        
        elif self._child_windows == False:
            print('mousePressEvent', event.x(), event.y(),)
            return QPushButton.mousePressEvent(self, event)

class PushButtonCombo(PushButton):
    

    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        
        #if self.color_ap == "":
        #    self._update_stylesheet(QColor(self.color_main), QColor("black"))
        #else:
        #    self._update_stylesheet(QColor(self.color_ap), QColor("black"))


    def _update_stylesheet(self, background, foreground):

        self.setStyleSheet(
            """
        QPushButton{
            background-color: %s;
            border: none;
            color: %s;
            text-align: center;
            text-decoration: none;
            font-size: %spx;
            border-radius: 2px;
            min-height: 14 px;
            padding: 2px 2px;
        }
        """
            #margin: 0.5px 0.5px;
            % (background.name(), foreground.name(), self._font_size)
        )

class PushButton_MainCombo(PushButton):
    

    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        
        #if self.color_ap == "":
        #    self._update_stylesheet(QColor(self.color_main), QColor("black"))
        #else:
        #    self._update_stylesheet(QColor(self.color_ap), QColor("black"))


    def _update_stylesheet(self, background, foreground):

        self.setStyleSheet(
            """
        QPushButton{
            background-color: %s;
            border: none;
            color: %s;
            padding: 4px 8px;
            text-align: center;
            text-decoration: none;
            font-size: %spx;
            border-radius: 8px;
            margin: 4px 2px;
            min-height: 20 px;
            border: 2px solid #B4B4B4;
        }
        """
            #margin: 0.5px 0.5px;
            % (background.name(), foreground.name(), self._font_size)
        )



class ComboBox(QComboBox):
    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        self.setStyleSheet("""
            QComboBox {
                border: 1px solid gray;
                border-radius: 3px;
                padding: 1px 18px 1px 3px;
                min-width: 6em;
            }


            QComboBox:drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;

                border-left-width: 1px;
                border-left-color: darkgray;
                border-left-style: solid; /* just a single line */
                border-top-right-radius: 3px; /* same radius as the QComboBox */
                border-bottom-right-radius: 3px;
            }

            QComboBox:down-arrow:on { /* shift the arrow when popup is open */
                top: 1px;
                left: 1px;
            }
            """)


class RadioButton(QRadioButton):

    def __init__(self, parent=None,font_size=None):
        super().__init__(parent)
        self._color = "#FFFFFF"
        self._color_new = "#B4B4B4"
        self._font_size = 14
        self._child_windows = False
        self._update_stylesheet(QColor(self._color),QColor("black"))
        self.setCursor(QCursor(Qt.PointingHandCursor))
        
        self._animation = QVariantAnimation(
            startValue=QColor(self._color_new),
            endValue=QColor(self._color),
            valueChanged=self._on_value_changed,
            duration=400,
        )

    
        
    def _update_stylesheet(self, background, foreground):
        self.setStyleSheet( """
            QRadioButton{
                background-color: %s;
                border: none;
                color: %s;
                padding: 4px 8px;
                text-align: center;
                text-decoration: none;
                font-size: %spx;
                border-radius: 8px;
                margin: 4px 2px;
                min-height: 20 px;
                }"""
                 % (background.name(), foreground.name(), self._font_size))
    
    @pyqtProperty(bool)
    def childWindows(self):
        return self._child_windows

    @childWindows.setter
    def childWindows(self, value):
        self._child_windows = value

    def _on_value_changed(self):
        foreground = (
            QColor("black")
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor("white")
        )
        background = (
            QColor(self._color)
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor(self._color_new)
            )
        self._update_stylesheet(background, foreground)

    def enterEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Backward)
        self._animation.start()
        super().enterEvent(event)

    def leaveEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Forward)
        self._animation.start()
        super().leaveEvent(event)

    def mousePressEvent(self, event):
        """
        if self._child_windows:
            msg = QMessageBox()
            msg.setWindowTitle("Информация")
            msg.setText("Закончите работу с открытым окном!!!")
            msg.setIcon(QMessageBox.Information)
            center = QPoint(self.geometry().center().x(), self.geometry().center().y())
            msg.show()
            msg.move(center)
            msg.exec_()
            print('!!!!!!!!!mousePressEvent!!!!!!!!!!!!!!!!!!', event.x(), event.y(),)
            return None
        elif self._child_windows == False:
            print('mousePressEvent', event.x(), event.y(),)
            return QRadioButton.mousePressEvent(self, event)
        """
        print('mousePressEvent', event.x(), event.y(),)
        return QRadioButton.mousePressEvent(self, event)



class Text_edit(QTextEdit):
    mouseDoubleClick = pyqtSignal()
    def __init__(self, parent=None,font_size=None):
        super().__init__(parent)
    def mouseDoubleClickEvent(self, event):
        print("pres IN TEXT EDIT", event)
        self.mouseDoubleClick.emit()
        return QTextEdit.mouseDoubleClickEvent(self,event)



class Common_status(QWidget):
    clear_form_signal = pyqtSignal()
    def __init__(self, width):
        super().__init__()
        self.width = width
        self.initUI()

    def initUI(self):
        
        self.font = QFont("Courier New", pointSize=12)
        self.all_objects = []
        self.line_objects = []
        self.localis=""
        self.localis_textedit = ""
        self.diagnosis =""
        self.diagnosis_textedit =""
        self.manipulation = ""
        self.manipulation_textedit = ""
        self.recomendation =""
        self.recomendation_textedit=""
        self.appointment= ""
        self.appointment_textedit = ""

        self.sample = ""

        print(datetime.date.today())


        if not self.layout():
            self.main_layout = QVBoxLayout()
            self.main_layout.setSpacing(0)
            self.main_layout.setContentsMargins(0,0,0,0)
        line_layout = QHBoxLayout()
        
        #self.common_status.setMaximumHeight(self.widget.height())
        #self.common_status.setMaximumWidth(self.widget.width())
        main_line = 1
        main_row = 1
        
        for obj in common_status_notation:
            if "date" in obj:
                date_widget = QDateEdit()
                date_widget.setDate(datetime.date.today())
                date_widget.setStyleSheet('''font-family: Courier New, monospace; font-size:16px;''')
                line_layout.addWidget(date_widget)
                date_widget.adjustSize()
                name = "date_widget"
                self.object_naming(text = obj, name = name,obj=date_widget, line=main_line, row=main_row)
                #date_widget.setObjectName("date_widget_"+str(main_line)+"_"+str(main_row))
                self.line_objects.append(date_widget)
                main_row += 1
                
            elif "textedit" in obj:

                #textedit = QTextEdit()
                textedit = Text_edit()
                
                name = "textedit"
                
                if "(" in obj:
                    text_1 = obj[obj.find("(")+1:obj.find(")")]
                    textedit.setText(text_1)
                    self.object_naming(text = obj[:obj.find("(")], name = name,obj=textedit, line=main_line, row=main_row)
                else:
                    self.object_naming(text = obj, name = name,obj=textedit, line=main_line, row=main_row)
                    
                #textedit.setMaximumHeight(25)
                textedit.setWordWrapMode(1)
                textedit.setAlignment(Qt.AlignVCenter)
                #textedit.setCurrentFont(self.font)
                textedit.setStyleSheet('''
                    font-family: Courier New, monospace;
                    font-size:16px;
                    ''')


                if textedit.toPlainText() != "":
                    #font = textedit.document().defaultFont()
                    fontMetrics = QFontMetrics(self.font)
                    textSize = fontMetrics.size(0, textedit.toPlainText())
                    w = textSize.width() + 25
                    h = textSize.height() +15
                    textedit.setMinimumSize(w, h)
                    textedit.setMaximumSize(w, h)
                    textedit.resize(w, h)

                else:
                    self.textedit_res_alternative()
                    


                #textedit.setSizeAdjustPolicy(0)
                #textedit.setWordWrapMode(1)
                textedit.textChanged.connect(self.textedit_res_alternative)

                line_layout.addWidget(textedit)
                #textedit.adjustSize()

                #textedit.setObjectName("textedit_"+str(main_line)+"_"+str(main_row))
                main_row += 1
                self.line_objects.append(textedit)
                
                
            elif "combobox" in obj:
                combobox = ComboBox()
                combobox.setSizeAdjustPolicy(0)
                #combobox.
                text_2 = obj[obj.find("[")+1:obj.find("]")]
                text_list = text_2.split(",")
                text_list_normalized = []
                for item in text_list:
                    text_list_normalized.append(item.strip())
                
                combobox.addItems(text_list_normalized)
                #combobox.setWordWrap(True)
                #combobox.setCurrentIndex(0)
                if "tooltip" in obj:
                    combobox.setToolTip(obj[obj.find("{")+1:obj.find("}")])
                combobox.setStyleSheet('''
                    font-family: Courier New, monospace;
                    font-size:16px;
                    ''')
                line_layout.addWidget(combobox)
                combobox.adjustSize()
                name = "combobox"
                self.object_naming(text = obj[:obj.find("(")], name = name,obj=combobox, line=main_line, row=main_row)
                #combobox.setObjectName("combobox_"+str(main_line)+"_"+str(main_row))
                main_row += 1
                self.line_objects.append(combobox)
                if "opt" in combobox.objectName():
                    combobox.activated.connect(self.combo_proccess)
                
            elif "lineedit" in obj:
                linedit = QLineEdit()
                
                if "(" in obj:
                    text_3 = obj[obj.find("(")+1:-1]
                    linedit.setText(text_3)
                line_layout.addWidget(linedit)
                
                linedit.setStyleSheet('''
                    font-family: Courier New, monospace;
                    font-size:16px;
                    ''')
                name = "linedit"
                if linedit.text() != "":
                    #font = textedit.document().defaultFont()
                    fontMetrics = QFontMetrics(self.font)
                    textSize = fontMetrics.size(0, linedit.text())
                    w = textSize.width() + 25
                    h = textSize.height() +15
                    linedit.setMinimumSize(w, h)
                    linedit.setMaximumSize(w, h)
                    linedit.resize(w, h)
                self.object_naming(text = obj[:obj.find("(")], name = name,obj=linedit, line=main_line, row=main_row)
                #linedit.setObjectName("linedit_"+str(main_line)+"_"+str(main_row))
                main_row += 1
                self.line_objects.append(linedit)
                
            elif "\n" in obj:
                widget = QWidget()

                if "opt" in obj:
                    text_name = obj[obj.rfind("_")+1:]
                    print("endline_text", text_name)
                    end_line = QLabel(text_name[:text_name.find("\n")])
                else:
                    end_line = QLabel(obj[:obj.find("\n")])
                
                end_line.setStyleSheet('''
                    font-family: Courier New, monospace;
                    font-size:16px;
                    margin-top:2px;
                    margin-bottom:2px;
                    padding-top:2px;
                    ''')
                line_layout.addWidget(end_line)
                name = "endline_label"
                self.object_naming(text = obj[:obj.find("\n")], name = name,obj=end_line, line=main_line, row=main_row)
                #end_line.setObjectName("endline_label_"+str(main_line)+"_"+str(main_row))
                main_row = 1
                main_line+=1
                self.line_objects.append(end_line)
                self.all_objects.append(self.line_objects)
                line_layout.addStretch(1)
                #widget.setMaximumWidth(self.width)
                
                widget.setLayout(line_layout)
                
                self.main_layout.addWidget(widget)
                self.line_objects = []
                line_layout = QHBoxLayout()


            else:
                if obj.startswith("opt"):
                    text_name= obj[obj.rfind("_")+1:]
                    lab = QLabel(text_name)
                    #lab.setWordWrap(True)
                    lab.hide()
                else:
                    lab = QLabel(obj)
                    #lab.setWordWrap(True)
                
                lab.setStyleSheet('''
                    font-family: Courier New, monospace;
                    font-size:16px;
                    margin-top:0ex;
                    margin-bottom:0ex;
                    ''')
                line_layout.addWidget(lab)
                name = "label"
                self.object_naming(text = obj, name = name,obj=lab, line=main_line, row=main_row)
                #lab.setObjectName("label_"+str(main_line)+"_"+str(main_row))
                main_row += 1
                self.line_objects.append(lab)
        self.main_layout.insertSpacing(-1, 30)
        but_layout = QHBoxLayout()
        button = PushButton("Сохранить в файле .txt")
        button.clicked.connect(self.text_proceess)
        button.color = "#ffb4a2"
        button.font_size = 14

        button_save = PushButton("Сохранить в файле .docx")
        button_save.clicked.connect(self.text_docx_save)
        button_save.color = "#ffb4a2"
        button_save.font_size = 14


        button_clear = PushButton("Очистить форму")
        button_clear.clicked.connect(self.clear_buttom_push)
        button_clear.color = '#fa8090'
        button_clear.font_size = 14
        but_layout.addWidget(button)
        but_layout.addWidget(button_save)
        but_layout.addWidget(button_clear)
        
        #self.main_layout.addLayout(but_layout)

        self.main_layout.addStretch(1)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        widget_upper = QWidget()
        widget_upper.setLayout(self.main_layout)
        scroll.setWidget(widget_upper)

        widget_lower = QWidget()
        widget_lower.setLayout(but_layout)
        

        self.common_layout = QVBoxLayout()
        self.common_layout.addWidget(scroll)
        self.common_layout.addWidget(widget_lower)


        if not self.layout():
            self.setLayout(self.common_layout)



    def object_naming(self, text, name,obj,line,row):
        if "opt" in text:
            if text.startswith("opt") and not text[:-1].endswith("opt"):
                obj.setObjectName("opt"+text[3]+"_"+name+"_"+str(line)+"_"+str(row))
                obj.hide()
            elif text.startswith("opt") and text[:-1].endswith("opt"):
                obj.setObjectName("opt"+text[3]+ "_"+name+"_"+"opt"+text[-1]+ "_"+str(line)+"_"+str(row))
                obj.hide()
            elif text[:-1].endswith("opt"):
                obj.setObjectName(name+"_"+"opt"+text[-1]+ "_"+str(line)+"_"+str(row))
        elif "localis" in text:
            obj.setObjectName(name+"localis_"+str(line)+"_"+str(row))
            self.localis_textedit = obj
        elif "manipulation" in text:
            obj.setObjectName(name+"manipulation_"+str(line)+"_"+str(row))
            obj.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=obj.toPlainText()))
            self.manipulation_textedit = obj
        elif "diagnosis" in text:
            obj.setObjectName(name+"diagnosis_"+str(line)+"_"+str(row))
            obj.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=obj.toPlainText()))
            self.diagnosis_textedit = obj
        elif "recomendation" in text:
            obj.setObjectName(name+"recomendation_"+str(line)+"_"+str(row))
            obj.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=obj.toPlainText()))
            self.recomendation_textedit = obj
        elif "appointment" in text:
            obj.setObjectName(name+"appointment_"+str(line)+"_"+str(row))
            obj.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=obj.toPlainText()))
            self.appointment_textedit = obj
        else:
            obj.setObjectName(name+"_"+str(line)+"_"+str(row))
        #print(text)
        #print(obj.objectName())
    
    def combo_proccess(self):
        combo = self.sender()
        combo_name = combo.objectName()
        options = [["opt1", "не пальпируется"], ["opt2", "не выступает"], ["opt3", "не пальпируется"]]
        for op in options:
            if op[0] in combo_name and combo.currentText() != op[1]:
                for line in self.all_objects:
                    for item in line:
                        if item.objectName().startswith(op[0]):
                            item.show()
            elif op[0] in combo_name and combo.currentText() == op[1]:
                for line in self.all_objects:
                    for item in line:
                        if op[0]=="opt1" and combo.currentText()=="не пальпируется":
                            if item.objectName().startswith(op[0]) or item.objectName().startswith("opt2"):
                                if "combobox" in item.objectName() and item.currentText()=="выступает":
                                    item.setCurrentIndex(0)

                                item.hide()
                        else:
                            if item.objectName().startswith(op[0]):
                                item.hide()

    def text_proceess(self):
        

        line_lenght = 70
        text_main = ""
        text_line = ""
        line_dict = []
        for line in self.all_objects:
            for item in line:
                if item.isVisible():
                    if "textedit" in item.objectName() and item.isVisible():
                        line_dict.append(item.toPlainText().strip() + " ")
                        text_line += item.toPlainText().strip() + " "
                        #print(item.toPlainText())
                        #print(item.isVisible())
                    elif "combobox" in item.objectName() and item.isVisible():
                        line_dict.append(item.currentText() + " ")
                        text_line += item.currentText() + " "
                        #print(item.currentText())
                        #print(item.isVisible())
                    elif "lineedit" in item.objectName() and item.isVisible():
                        line_dict.append(item.toPlainText().strip() + " ")
                        text_line += item.toPlainText().strip() + " "
                    
                    else:
                        if item.text().strip() in [",", ".", ":", ";"]:
                            print("item text strip()")
                            print(line_dict[-1])
                            text_line = text_line.rstrip() +item.text().strip() +" "
                            line_dict[-1] = line_dict[-1].rstrip()
                            line_dict.append(item.text().lstrip() + " ")
                        else:
                            line_dict.append(item.text().strip() + " ")
                            text_line += item.text().strip() + " "
                        #print(item.text())
                        #print(item.isVisible())
            print("text_line", text_line)
            print("line_dict", line_dict)

            if len(text_line)>line_lenght:
                print("text_line", text_line)
                print("line_dict", line_dict)

                formated_line = ""
                for element in line_dict:
                    if element.strip() in [",", ".", ":", ";"]:
                        formated_line = formated_line.rstrip() +element.lstrip()
                    else:
                        formated_line+=element
                    #formated_line = formated_line.strip()


                    if len(formated_line)>line_lenght:
                        print("formated_line", formated_line)
                        line_to_add = formated_line[:line_lenght]
                        formated_line = formated_line[line_lenght:]
                        while line_to_add[-1] !=" ":
                            formated_line=line_to_add[-1] + formated_line
                            line_to_add=line_to_add[:-1]
                            print(line_to_add)
                        
                        line_to_add = line_to_add.strip()



                        print("line_to_add", line_to_add)
                        delta_lenght = line_lenght - len(line_to_add)
                        words_in_line_dict = line_to_add.split(" ")
                        spaces_in_line = len(words_in_line_dict) - 1
                        spaces_map = []
                        print(delta_lenght)
                        print(spaces_in_line)
                        space_lenght = delta_lenght//spaces_in_line
                        space_extra = delta_lenght%spaces_in_line
                        for i in range(spaces_in_line):
                            if space_extra != 0:
                                spaces_map.append(space_lenght+1)
                                space_extra-1
                            else:
                                spaces_map.append(space_lenght)

                        line_new = ""
                        increment = 0
                        for word in words_in_line_dict:
                            if increment<=len(spaces_map)-1:
                                line_new+=word + " "*spaces_map[increment]
                            else:
                                line_new+=word
                            increment+=1

                        text_main+=line_new +"\n"
                print(formated_line)
                text_main += formated_line + "\n"
                formated_line = ""
                line_dict = []
                text_line = ""
            else:
                text_main+=text_line +"\n"
                text_line = ""
                line_dict = []
            
            print("new_line")


        desktop_destination = os.path.join(os.environ['USERPROFILE'], "Desktop\\ORTHO_STATUS\\")
        print(desktop_destination)
        try:
            os.mkdir(desktop_destination, mode=stat.S_ISUID)
        except OSError as error:
            print(error)   
        
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self,"Сохранение файла .txt",directory=desktop_destination, 
                                                    filter="All Files (*);;Text Files (*.txt)", options=options)
        if fileName:
            print(fileName)
        if fileName[-4:] !=".txt":
            full_way = fileName +".txt"
        else:
            full_way = fileName
        print(full_way)


        with open(full_way, "w") as f:
            f.write("Дата %s, время: %s.\n\n" %(time.strftime("%d.%m.%Y"), time.strftime("%H:%M")))
            for line in text_main:
                f.write(line)
            f.close()

    def text_docx_save(self):
        mydoc = docx.Document()
        current_section = mydoc.sections[-1]
        current_section.right_margin = Cm(1.0)
        current_section.top_margin = Cm(1.0)
        current_section.bottom_margin = Cm(1.0)
        par = mydoc.add_paragraph(style="No Spacing")
        for line in self.all_objects:
            
            for item in line:
                if item.isVisible():
                    if "textedit" in item.objectName():
                        #print(item.objectName())

                        run = par.add_run(item.toPlainText())
                        font = run.font
                        font.name = "Courier New"
                        font.size = Pt(10)
                        #print(item.toPlainText())
                    elif "linedit" in item.objectName():
                        run = par.add_run(item.text())
                        font = run.font
                        font.name = "Courier New"
                        font.size = Pt(10)

                    elif "combobox" in item.objectName():
                        run = par.add_run(item.currentText())
                        font = run.font
                        font.name = "Courier New"
                        font.size = Pt(10)
                    else:
                        text = item.text()[:]
                        if "\n" in text: 
                            text = text.replace("\n", "")
                        #print(text)
                        if text in (["Жалобы: ", "Анамнез: ", "Данные объективного обследования", 
                                    "Локальный статус: ", "Проводимые манипуляции: ", "Диагноз: ", 
                                    "Назначенные обследования: ", "Рекомендации: ", "Врач ", "Группа 'Д' наблюдения "]):
                            par.add_run("\n")
                            run = par.add_run(text)
                            
                            if text not in ["Врач ", "Группа 'Д' наблюдения "]:
                                
                                run.bold = True
                                run.underline = True
                            elif text == "Врач ":
                                par.runs[-2].text = "\n\n"

                            font = run.font
                            font.name = "Courier New"
                            font.size = Pt(10)
                        else:

                            run = par.add_run(text)
                            font = run.font
                            font.name = "Courier New"
                            font.size = Pt(10)
                        #print(text)

                    if not par.runs[-1].text.endswith(" "):
                        run = par.add_run(" ")
                        font = run.font
                        font.name = "Courier New"
                        font.size = Pt(10)
        
        for iteration in range(3):
            previous_run_text = ""
            previous_run = ""
            for r in par.runs:
                if r.text !="":
                    #print("before normalization:")
                    #print(r.text)
                    r.text = r.text.replace("  ", " ")
                    r.text = r.text.replace("   ", " ")
                    r.text = r.text.replace("    ", " ")
                    r.text = r.text.replace(" ,", ",")
                    r.text = r.text.replace(" .", ".")
                    r.text = r.text.replace(" :", ":")
                    r.text = r.text.replace(" ;", ";")
                    r.text = r.text.replace("..", ".")
                    r.text = r.text.replace(",,", ",")
                    #print(r.text)
                    #print("previous_run_text")
                    #print(previous_run_text)
                    #print("len", len(previous_run_text))
                    if previous_run_text !="":
                        if previous_run_text[-1] == " " and r.text[0] ==" ": r.text = r.text[1:]
                        if r.text != "":
                            if previous_run_text[-1] == " " and r.text[0] =="." and previous_run !="": 
                                previous_run.text = previous_run_text[:-1]
                                previous_run_text = previous_run.text[:]
                            if previous_run_text != "":
                                if len(r.text) == 1 and len(previous_run_text) == 1:
                                    if previous_run_text == "." and r.text ==".": r.text =""
                                elif len(r.text) == 1 and len(previous_run_text) > 1:
                                    if previous_run_text[-1] == "." and r.text ==".": r.text =""
                                elif len(r.text) > 1 and len(previous_run_text) == 1:
                                    if previous_run_text == "." and r.text[0] ==".": r.text =r.text[1:]
                                elif len(r.text) > 1 and len(previous_run_text) != 1:
                                    if previous_run_text[-1] == "." and r.text[0] ==".": r.text = r.text[1:]
                            
                    previous_run_text = r.text[:]
                    previous_run = r
                    #print("after  normalization:")
                    #print(r.text)
                else:
                    continue

        desktop_destination = os.path.join(os.environ['USERPROFILE'], "Desktop\\ORTHO_STATUS\\")
        print(desktop_destination)
        try:
            os.mkdir(desktop_destination, mode=stat.S_ISUID)
        except OSError as error:
            print(error)

        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self,"Сохранение файла .docx",directory=desktop_destination, 
                                                    filter="All Files (*);;Text Files (*.docx)", options=options)
        if fileName:
            print(fileName)
        if fileName[-5:] !=".docx":
            full_way = fileName +".docx"
        else:
            full_way = fileName
        print(full_way)
        
        mydoc.save(full_way)
    
    def text_edit_double_click_process(self, text):
        obj = self.sender()
        print(obj.objectName())
        with open('samples_manipulation.pickle',  'rb') as f:
            manipulation  = pickle.load(f)

        with open('samples_diagnosis.pickle',  'rb') as f:
            diagnosis  = pickle.load(f)
        
        with open('samples_appointment.pickle', 'rb') as z:
            appointment= pickle.load(z)
        
        with open('samples_recomendation.pickle', 'rb') as x:
            recomendation= pickle.load(x)
        


        if "diagnosis" in obj.objectName():
            self.sample = SampleWindow (source=diagnosis, source_name='samples_diagnosis.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
        elif "manipulation" in obj.objectName():
            self.sample = SampleWindow (source=manipulation, source_name='samples_manipulation.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
        elif "appointment" in obj.objectName():
            self.sample = SampleWindow (source=appointment, source_name='samples_appointment.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
        elif "recomendation" in obj.objectName():
            self.sample = SampleWindow (source=recomendation, source_name='samples_recomendation.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
    
    def text_pasting(self, obj, text):
        if text !="":
            obj.setText(text)
            self.textedit_res_alternative()


    def textedit_res_alternative(self):
        

        textedit_to_resize = self.sender()
        print("self.width=", self.width)
        max_width = self.width
        #sizes = []
        
        for line in self.all_objects:
            line_width = 0
            textedit_in_line = 0
            textedit_in_line_obj = []
            for it in line:
                font_other = QFont("Courier New", pointSize=12)
                fontMetrics_other = QFontMetrics(font_other)
                if "label" in it.objectName() or "lineedit" in it.objectName():
                    other_textSize = fontMetrics_other.size(0, it.text())
                elif "textedit" in it.objectName():
                    #other_textSize = fontMetrics_other.size(0, it.toPlainText())
                    
                    if "diagnosis" in it.objectName():
                        self.diagnosis = it.toPlainText()
                        self.diagnosis_textedit = it
                    elif "manipulation" in it.objectName():
                        self.manipulation = it.toPlainText()
                        self.manipulation_textedit = it
                    elif "recomendation" in it.objectName():
                        self.recomendation = it.toPlainText()
                        self.recomendation_textedit = it
                    elif "localis" in it.objectName():
                        self.localis = it.toPlainText()
                        self.localis_textedit = it
                    elif "appointment" in it.objectName():
                        self.appointment = it.toPlainText()
                        self.appointment_textedit = it


                    textedit_in_line +=1
                    textedit_in_line_obj.append(it)
                    continue
                elif "combobox" in it.objectName():
                    other_textSize = fontMetrics_other.size(0, it.currentText())
                else:
                    other_textSize = fontMetrics_other.size(0, it.text())
                line_width += other_textSize.width()

                

            if textedit_in_line >0:
                for textedit in textedit_in_line_obj:
                    font_other = QFont("Courier New", pointSize=12)
                    fontMetrics_other = QFontMetrics(font_other)
                    other_textSize = fontMetrics_other.size(0, textedit.toPlainText())
                    text_width = other_textSize.width()
                    w= (max_width-line_width)/textedit_in_line
                    if textedit.toPlainText()=="" or text_width>w:

                        if textedit.toPlainText()=="":
                            textedit.setMinimumWidth(w-30)
                            textedit.setMaximumWidth(w-30)
                            textedit.setMinimumHeight(50)
                            textedit.setMaximumHeight(30*3)
                            textedit.resize(w-30,50)
                            print("texted without text",textedit.size())
                        else:
                            textedit.setMinimumWidth(w-25)
                            textedit.setMaximumWidth(w-25)
                            textedit.setMinimumHeight(30*3)
                            textedit.setMaximumHeight(30*3)
                            textedit.resize(w-25,30*3)
                            print("texted wide text",textedit.size())
                        print(textedit.objectName(), textedit.toPlainText())
                    else:
                        if (textedit.toPlainText() in 
                            ["нет", "не обнаружено", "отказался", ". ", "не пигментированы", "отрицательные",
                            "отсутствует.", "безболезненный", "не учащен", "норма", "травматолог-ортопед"]):
                            textedit.setMinimumWidth(text_width+15)
                            textedit.setMaximumWidth(text_width+15)
                            textedit.resize(text_width+15, other_textSize.height()+20)
                        else:
                            textedit.setMinimumWidth(w-20)
                            textedit.setMaximumWidth(w-20)
                            textedit.resize(w-20, other_textSize.height()+20)
                        textedit.setMinimumHeight(other_textSize.height()+20)
                        textedit.setMaximumHeight(other_textSize.height()+20)
                        print("texted normal text width",textedit.size())
                    #print(textedit.objectName(), textedit.toPlainText())
            #print("line font width", line_width)

    '''
        for line in self.all_objects:
            for obj in line:
                if obj.objectName() == textedit_to_resize.objectName():
                    other_wid_width = 0
                    print(obj.objectName(), obj.toPlainText())
                    for i in line:
                        
                        if i.objectName()!=obj.objectName():
                            
                            #font_other = obj.document().defaultFont()
                            #fontMetrics_other = QFontMetrics(font_other)
                            font_other = QFont("Courier New", pointSize=12)
                            fontMetrics_other = QFontMetrics(font_other)

                            if "label" in i.objectName() or "lineedit" in i.objectName():
                                other_textSize = fontMetrics_other.size(0, i.text())
                            elif "textedit" in i.objectName():
                                other_textSize = fontMetrics_other.size(0, i.toPlainText())
                            elif "combobox" in i.objectName():
                                other_textSize = fontMetrics_other.size(0, i.currentText())
                            
                            other_wid_width += other_textSize.width()
                            print(i.objectName())
                            print(other_wid_width)
                    
                    #font = obj.document().defaultFont()
                    if obj.toPlainText() != "":
                        font = QFont("Courier New", pointSize=11)
                        fontMetrics = QFontMetrics(font)

                        textSize = fontMetrics.size(1, obj.toPlainText())
                        #textSize.maxWidth()
                        w = textSize.width()
                        h = textSize.height()
                        
                        print(w,h)
                        

                        #w = textSize.width() + 10
                        #h = textSize.height() +10

                        if w > max_width - other_wid_width:
                            w = max_width - other_wid_width
                        
                        #obj.resize(w, h)
                        textSize_after = fontMetrics.size(1, obj.toPlainText())
                        print(textSize_after.width(),textSize_after.height())

                        #    w = self.width - other_wid_width
               
                        #obj.setMinimumWidth(w)
                        #obj.setMaximumWidth(w)
                        #obj.resize(w, h)

                        print(obj.objectName(), obj.size())
                        
                        #print(w, h)
                        #print(self.width)

                    if "diagnosis" in obj.objectName():
                        self.diagnosis = obj.toPlainText()
                        self.diagnosis_textedit = obj
                    elif "manipulation" in obj.objectName():
                        self.manipulation = obj.toPlainText()
                        self.manipulation_textedit = obj
                    elif "recomendation" in obj.objectName():
                        self.recomendation = obj.toPlainText()
                        self.recomendation_textedit = obj
                    elif "localis" in obj.objectName():
                        self.localis = obj.toPlainText()
                        self.localis_textedit = obj
                    elif "appointment" in obj.objectName():
                        self.appointment = obj.toPlainText()
                        self.appointment_textedit = obj
    '''


    def setLocalis(self, text):
        print("in set localis")
        for line in self.all_objects:
            
            for item in line:
                #print(item.objectName())
                #print("item width", item.width())
                #print("item height", item.height())
                if "localis" in item.objectName():
                    item.setText(text)
                    self.localis=text
                    self.textedit_res_alternative()
                    #item.textChanged.emit()

    def clear_buttom_push(self):
        self.clear_form_signal.emit()
        self.clear_form()


    def clear_form(self):
        
        layout = self.layout()
        print(layout)
        self.clearLayout(layout)
        import sip
        sip.delete(layout)
        self.initUI()
        #for line in self.all_objects:
        #    for item in line:
        #        if "textedit" in item.objectName():
        #            item.setText()

    def clearLayout(self, layout):
        if layout != None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget() is not None:
                    child.widget().deleteLater()
                elif child.layout() is not None:
                    self.clearLayout(child.layout())

class Label_press(QLabel):
    label_pressed = pyqtSignal()
    label_double_pressed = pyqtSignal()
    def __init__(self):
        super().__init__()
    
    def mousePressEvent(self,event):
        self.label_pressed.emit()
        print("press")

        return QLabel.mousePressEvent(self,event)



class SampleWindow(QWidget):
    sample_window_text_sending = pyqtSignal()
    def __init__(self,source, source_name, text, parent=None):
        super().__init__()
        self.font = QFont("Courier New", pointSize=14)
        self.source = source
        self.source_name = source_name
        self.text = text
        self.initUI()

    def initUI(self):
        
        self.add_group_box = ""

        main_layout = QGridLayout(self)
        main_layout.addWidget(QLabel("Текстовое поле"), 0,0)
        main_layout.addWidget(QLabel("Шаблоны"), 0, 1)

        self.left_box = QTextEdit()
        self.left_box.setText(self.text)
        main_layout.addWidget(self.left_box, 1,0,1,1)
        self.right_box = QVBoxLayout()
        self.widget = QWidget()
        self.widget.setLayout(self.right_box)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(self.widget)
        main_layout.addWidget(scroll,1,1,1,1)
        #but_edit = PushButton("Редактировать список")
        #but_edit.clicked.connect(self.table_display)
        #but_edit.color = "#ffb4a2"
        #but_edit.font_size = 14

        but_ok = PushButton("Ok")
        but_ok.clicked.connect(self.close_sample_window)
        but_ok.color = "#adff2f"
        but_ok.font_size = 14

        #main_layout.addWidget(but_edit,2,1,1,1)
        main_layout.addWidget(but_ok,2,0,1,2)
        buttons= QHBoxLayout()
        self.resize(640,480)

        width = self.widget.width()
        print(self.width())
        
        self.labels_samples = []
        self.but_ok_list = []
        #print("width", width)
        
        self.building_element_tree()

        self.add_editing_chapter_key()
        




        self.show()
    

    def building_element_tree(self):
        self.clearLayout(self.right_box)
        self.labels_samples= []
        self.but_ok_list=[]
        
        for key, value in self.source.items():
            group_box = QGroupBox(key)
            group_box.setStyleSheet('''
                QGroupBox {
                    margin-top: 2ex;
                    font-size: 14px;
                    
                }
                QGroupBox:enabled {
                    border: 3px solid #CA8372;
                    border-radius: 5px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 3ex;
                }
            ''')
            if isinstance(value, list):
                but_layout = QVBoxLayout()
                but_layout.setSpacing(2)
                increment = 0
                for item in value:
                    line_box = QHBoxLayout()
                    press_lab = Label_press()
                    if len(item) > 40:
                        text_new = ""
                        inc = 1
                        for i in item:
                            text_new+=i
                            inc+=1
                            if inc ==41:
                                text_new+="\n"
                                inc = 1
                        press_lab.setText(text_new)
                    else:
                        press_lab.setText(item)
                    self.labels_samples.append(press_lab)
                    press_lab.setObjectName("lab_" + key +"}"+ str(increment))
                    press_lab.setStyleSheet('''border: 1px solid #CA8372; padding: 4 px; border-radius: 3px''')
                    press_lab.label_pressed.connect(self.label_pressed_procees)
                    #button.setMaximumWidth(100)
                    edit_layout = QVBoxLayout()
                    but_edit = QPushButton("Edit")
                    but_edit_ok = QPushButton("Ok")
                    self.but_ok_list.append(but_edit_ok)
                    but_edit.clicked.connect(self.label_editing)
                    #but_edit_ok.clicked.connect(self.but_ok_proccess)
                    but_edit.setObjectName("but_" + key +"}"+ str(increment))
                    but_edit_ok.setObjectName("but_ok_" + key  +"}" + str(increment))
                    line_box.addWidget(press_lab)
                    edit_layout.addWidget(but_edit)
                    edit_layout.addWidget(but_edit_ok)
                    line_box.addLayout(edit_layout)
                    but_layout.addLayout(line_box)
                    

                    but_edit_ok.hide()
                    print(press_lab.width())
                    increment+=1
                but_add = QPushButton("Добавить элемент")
                but_add.setObjectName("but_add_" + key)
                but_add.clicked.connect(self.adding_press_lab)
                but_layout.addWidget(but_add)
            
            group_box.setLayout(but_layout)
            self.right_box.addWidget(group_box)




    def adding_press_lab(self):
        text, ok = QInputDialog.getText(self, 'Input Dialog', 'Enter text:')
        if ok:
            self.source[self.sender().objectName()[8:]].append(text)
            with open(self.source_name, 'wb') as f:
                    pickle.dump(self.source, f)
            self.building_element_tree()
            self.add_editing_chapter_key()



    def add_editing_chapter_key(self):
        if self.add_group_box != "":
            self.add_group_box.setParent(None)
            self.add_group_box = ""
        
        self.add_group_box = QGroupBox("Редактировать раздел")
        self.add_group_box.setStyleSheet('''
                QGroupBox {
                    margin-top: 2ex;
                    font-size: 14px;
                    
                }
                QGroupBox:enabled {
                    border: 3px solid #CA8372;
                    border-radius: 5px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 3ex;
                }
            ''')
        self.button_add = QPushButton("Довавить")
        self.button_add.setObjectName("add")
        self.button_del_chap = QPushButton("Удалить раздел")
        self.button_del_chap.setObjectName("del")
        self.button_del_chap.clicked.connect(self.chapter_edit_proccess)
        self.button_add.clicked.connect(self.chapter_edit_proccess)
        

        but_lay_chapter = QHBoxLayout()
        but_lay_chapter.addWidget(self.button_add)
        but_lay_chapter.addWidget(self.button_del_chap)
        self.add_group_box.setLayout(but_lay_chapter)

        self.right_box.addWidget(self.add_group_box)




    def label_editing(self):
        but_name = self.sender().objectName()
        label = ""
        text = ""
        key = but_name[4:but_name.find("}")]
        label_index = int(but_name[but_name.find("}")+1:])
        print(label_index)
        for lab in self.labels_samples:
            print(lab.objectName())
            if lab.objectName()[4:] == but_name[4:]:
                print(lab.objectName())
                label = lab
                text = lab.text()
        print(text)
        print(key)
        text_input, ok = QInputDialog.getMultiLineText(self, 'Редактирование элемента', 'Откорректируйте текст', text)
        if ok and text_input!="":
            text_input = text_input.replace("\n", "")
            
            self.source[key][label_index] = text_input 
            with open(self.source_name, 'wb') as f:
                pickle.dump(self.source, f)
            self.building_element_tree()
            self.add_editing_chapter_key()

        elif ok and text_input =="":
            self.source[key].pop(label_index)
            with open(self.source_name, 'wb') as f:
                pickle.dump(self.source, f)
            self.building_element_tree()
            self.add_editing_chapter_key()


    

    def chapter_edit_proccess(self):
        if self.sender().objectName() == "add":
            text, ok = QInputDialog.getText(self, 'Input Dialog', 'Enter text:')
            if ok:
                self.source[text] = []
                with open(self.source_name, 'wb') as f:
                    pickle.dump(self.source, f)
                self.building_element_tree()
                self.add_editing_chapter_key()

        elif self.sender().objectName() == "del":
            if self.source.keys() != []:
                item, ok = QInputDialog.getItem(self, 'Input Dialog', 'Enter text:', self.source.keys(), 0)
                if ok:
                    self.source.pop(item)
                    with open(self.source_name, 'wb') as f:
                        pickle.dump(self.source, f)
            else:

                msg = QMessageBox()
                msg.setWindowTitle("Информация")
                msg.setText("Информация не введена!!!")
                msg.setIcon(QMessageBox.Information)
                msg.show()
                msg.exec_()


    def editing_proccess(self, lab, texted):
        print(texted.toPlainText())
        lab.setText(texted.toPlainText())
        texted.setParent(None)
        self.sender().hide()



    def table_display(self):
        self.table = Table_view(self.source)
        self.table.show()

    def close_sample_window(self):
        self.sample_window_text_sending.emit()
        self.close()

    def label_pressed_procees(self):
        self.sender().text()
        text_cleaned = self.sender().text().replace("\n","")
        print("text_cleaned=", text_cleaned)
        print(self.left_box.toPlainText())
        if self.left_box.toPlainText() != "":
            self.left_box.setText(self.left_box.toPlainText() + "\n" + text_cleaned)
        else:
            self.left_box.setText(self.left_box.toPlainText() +text_cleaned)

    def clearLayout(self, layout):
        if layout != None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget() is not None:
                    child.widget().deleteLater()
                elif child.layout() is not None:
                    self.clearLayout(child.layout())

class Table_view(QTableWidget):
    def __init__(self, data, parent=None):
        QTableWidget.__init__(self, parent=None)
        self.data = data
        self.setData()
        self.resizeColumnsToContents()
        self.resizeRowsToContents()

    def setData(self): 
        horHeaders = []
        print(self.data.keys())
        for n, key in enumerate(sorted(self.data.keys())):
            header_item = QTableWidgetItem(key)
            self.setItem(n, 1, header_item)
            horHeaders.append(key)
            for m, item in enumerate(self.data[key]):
                for z, i in enumerate(item):
                    newitem = QTableWidgetItem(i)
                    self.setItem(m, z, newitem)
            #empty_item = QTableWidgetItem()
            #self.setItem(len(self.data[key])+1, n, empty_item)
        self.setHorizontalHeaderLabels(horHeaders)
        print("table", self)



