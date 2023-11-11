#!/usr/bin/python3
# -*- coding: utf-8 -*-

import sys
import os
import pickle
import logging
import datetime
import time
import stat
import docx

from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from picture_blank_making import add_float_picture

from PyQt5.QtWidgets import (QWidget, QTextEdit, QLineEdit, QPushButton, QScrollArea, QScrollBar, QSizePolicy,
    QHBoxLayout, QVBoxLayout, QApplication, QLabel, QRadioButton, QButtonGroup, QComboBox, QInputDialog, 
    QGridLayout, QGroupBox, QMessageBox, QDesktopWidget, QFileDialog, QTabWidget, QDateEdit, QListWidget)
from PyQt5.QtCore import QVariant, QEvent, QPoint
from PyQt5.QtCore import (Qt, pyqtSignal, QRect, QVariantAnimation,QEasingCurve, QAbstractAnimation, 
                        QPropertyAnimation, pyqtProperty)
from PyQt5.QtGui import QFont, QFontMetrics, QPainter ,QPixmap, QImage, QColor, QPen, QCursor, QTextCharFormat, QTextCursor, QIcon
from common_notation import common_status_notation



class PushButton(QPushButton):
    

    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        
        #if self.color_ap == "":
        #    self._update_stylesheet(QColor(self.color_main), QColor("black"))
        #else:
        #    self._update_stylesheet(QColor(self.color_ap), QColor("black"))
        self.parent = parent
        self._color = "#ffb4a2"
        self._color_new = "#A64A35"
        self._font_size = 14
        self._child_windows = False
        self.setCursor(QCursor(Qt.PointingHandCursor))
        
        self._animation = QVariantAnimation(
            startValue=QColor(self._color_new),
            endValue=QColor(self._color),
            valueChanged=self._on_value_changed,
            duration=400,
        )

    @pyqtProperty(str)
    def color(self):
        return self._color

    @color.setter
    def color(self, value):
        self._color = value
        color_new = ""
        if value == "#f0f0f0": 
            color_new = "#B4B4B4"
        
        if value == "#B4B4B4": 
            color_new = "#D2D2D2"

        elif value == "#adff2f":
            color_new = "#6AA60F"
        elif value =="#ffb4a2":
            color_new = "#A64A35"

        elif value =="#A64A35":
            color_new = "#ffb4a2"

        elif value =="":
            color_new = ""
        elif value == "#fa8090":
            color_new = "#BB7780"
        
        elif value == "#FF7373":
            color_new = "#BF3030"

        elif value == "#FF4040":
            color_new = "#BF3030"

        elif value == "#BF3030":
            color_new = "#FF7373"


        self._color_new = color_new
        self._update_stylesheet(QColor(self._color),QColor("black"))

    @pyqtProperty(str)
    def color_new(self):
        return self._color_new

    @color_new.setter
    def color_new(self, color):
        if color == "#f0f0f0": 
            color_new = "#B4B4B4"
        elif color == "#adff2f":
            color_new = "#91BF4A"
        elif color =="#ffb4a2":
            color_new = "#A64A35"
        elif color =="":
            color_new = ""
        self._color_new = color_new

    @pyqtProperty(bool)
    def childWindows(self):
        return self._child_windows

    @childWindows.setter
    def childWindows(self, value):
        self._child_windows = value

    def setColor (self, color):
        
        return color
        #self._update_stylesheet(QColor(self.color_main),QColor("black"))
        #self.update()

    @pyqtProperty(int)
    def font_size(self):
        return self._font_size

    @font_size.setter
    def font_size(self, font_size):
        self._font_size = font_size

    def getColor (self):
        return self.color_main

    def _on_value_changed(self):
        foreground = (
            QColor("black")
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor("white")
        )
        background = (
            QColor(self._color)
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor(self._color_new)
            )
        self._update_stylesheet(background, foreground)


    def _update_stylesheet(self, background, foreground):

        self.setStyleSheet(
            """
        QPushButton{
            background-color: %s;
            border: none;
            color: %s;
            padding: 4px 8px;
            text-align: center;
            text-decoration: none;
            font-size: %spx;
            border-radius: 8px;
            margin: 4px 2px;
            min-height: 20 px;
            
        }
        """
            % (background.name(), foreground.name(), self._font_size)
        )

    def enterEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Backward)
        self._animation.start()
        super().enterEvent(event)

    def leaveEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Forward)
        self._animation.start()
        super().leaveEvent(event)

    def mousePressEvent(self, event):
        
        if self._child_windows:
            msg = QMessageBox()
            msg.setWindowTitle("Информация")
            msg.setText("Закончите работу с открытым окном!!!")
            msg.setIcon(QMessageBox.Information)
            center = [QDesktopWidget().screenGeometry(-1).width()/2, QDesktopWidget().screenGeometry(-1).height()/2] 
            msg.show()
            msg.move(center[0],center[1])
            msg.exec_()
            print('!!!!!!!!!mousePressEvent!!!!!!!!!!!!!!!!!!', event.x(), event.y(),)
            return None
        
        elif self._child_windows == False:
            print('mousePressEvent', event.x(), event.y(),)
            return QPushButton.mousePressEvent(self, event)

class PushButtonCombo(PushButton):
    

    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        
        #if self.color_ap == "":
        #    self._update_stylesheet(QColor(self.color_main), QColor("black"))
        #else:
        #    self._update_stylesheet(QColor(self.color_ap), QColor("black"))


    def _update_stylesheet(self, background, foreground):

        self.setStyleSheet(
            """
        QPushButton{
            background-color: %s;
            border: none;
            color: %s;
            text-align: center;
            text-decoration: none;
            font-size: %spx;
            border-radius: 2px;
            min-height: 14 px;
            padding: 2px 2px;
        }
        """
            #margin: 0.5px 0.5px;
            % (background.name(), foreground.name(), self._font_size)
        )

class PushButton_MainCombo(PushButton):
    

    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        
        #if self.color_ap == "":
        #    self._update_stylesheet(QColor(self.color_main), QColor("black"))
        #else:
        #    self._update_stylesheet(QColor(self.color_ap), QColor("black"))


    def _update_stylesheet(self, background, foreground):

        self.setStyleSheet(
            """
        QPushButton{
            background-color: %s;
            border: none;
            color: %s;
            padding: 4px 8px;
            text-align: center;
            text-decoration: none;
            font-size: %spx;
            border-radius: 8px;
            margin: 4px 2px;
            min-height: 20 px;
            border: 2px solid #B4B4B4;
        }
        """
            #margin: 0.5px 0.5px;
            % (background.name(), foreground.name(), self._font_size)
        )



class ComboBox(QComboBox):
    def __init__(self, parent=None,font_size=None, child_windows=None):
        super().__init__(parent)
        self.setStyleSheet("""
            QComboBox {
                border: 1px solid gray;
                border-radius: 3px;
                padding: 1px 18px 1px 3px;
                min-width: 6em;
            }


            QComboBox:drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;

                border-left-width: 1px;
                border-left-color: darkgray;
                border-left-style: solid; /* just a single line */
                border-top-right-radius: 3px; /* same radius as the QComboBox */
                border-bottom-right-radius: 3px;
            }

            QComboBox:down-arrow:on { /* shift the arrow when popup is open */
                top: 1px;
                left: 1px;
            }
            """)


class RadioButton(QRadioButton):

    def __init__(self, parent=None,font_size=None):
        super().__init__(parent)
        self._color = "#FFFFFF"
        self._color_new = "#B4B4B4"
        self._font_size = 14
        self._child_windows = False
        self._update_stylesheet(QColor(self._color),QColor("black"))
        self.setCursor(QCursor(Qt.PointingHandCursor))
        
        self._animation = QVariantAnimation(
            startValue=QColor(self._color_new),
            endValue=QColor(self._color),
            valueChanged=self._on_value_changed,
            duration=400,
        )

    
        
    def _update_stylesheet(self, background, foreground):
        self.setStyleSheet( """
            QRadioButton{
                background-color: %s;
                border: none;
                color: %s;
                padding: 4px 8px;
                text-align: center;
                text-decoration: none;
                font-size: %spx;
                border-radius: 8px;
                margin: 4px 2px;
                min-height: 20 px;
                }"""
                 % (background.name(), foreground.name(), self._font_size))
    
    @pyqtProperty(bool)
    def childWindows(self):
        return self._child_windows

    @childWindows.setter
    def childWindows(self, value):
        self._child_windows = value

    def _on_value_changed(self):
        foreground = (
            QColor("black")
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor("white")
        )
        background = (
            QColor(self._color)
            if self._animation.direction() == QAbstractAnimation.Forward
            else QColor(self._color_new)
            )
        self._update_stylesheet(background, foreground)

    def enterEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Backward)
        self._animation.start()
        super().enterEvent(event)

    def leaveEvent(self, event):
        #if self.color_main == "#f0f0f0": 
        #    self.start_color = "#B4B4B4"
        #elif self.color_main == "#adff2f":
        #    self.start_color = "#91BF4A"
        
        self._animation.setDirection(QAbstractAnimation.Forward)
        self._animation.start()
        super().leaveEvent(event)

    def mousePressEvent(self, event):
        """
        if self._child_windows:
            msg = QMessageBox()
            msg.setWindowTitle("Информация")
            msg.setText("Закончите работу с открытым окном!!!")
            msg.setIcon(QMessageBox.Information)
            center = QPoint(self.geometry().center().x(), self.geometry().center().y())
            msg.show()
            msg.move(center)
            msg.exec_()
            print('!!!!!!!!!mousePressEvent!!!!!!!!!!!!!!!!!!', event.x(), event.y(),)
            return None
        elif self._child_windows == False:
            print('mousePressEvent', event.x(), event.y(),)
            return QRadioButton.mousePressEvent(self, event)
        """
        print('mousePressEvent', event.x(), event.y(),)
        return QRadioButton.mousePressEvent(self, event)

class Text_edit(QTextEdit):
    mouseDoubleClick = pyqtSignal()
    def __init__(self, parent=None,font_size=None):
        super().__init__(parent)
    def mouseDoubleClickEvent(self, event):
        print("pres IN TEXT EDIT", event)
        self.mouseDoubleClick.emit()
        return QTextEdit.mouseDoubleClickEvent(self,event)




class Consult(QWidget):
    clear_form_signal = pyqtSignal()
    def __init__(self, width):
        super().__init__()
        #self.width = width
        self.initUI()

    def initUI(self):
        self.font = QFont("Courier New", pointSize=12)
        self.all_objects = []
        self.line_objects = []
        self.localis = ""
        self.manipulation = ""
        self.diagnosis =""
        self.recomendation = ""
        self.appointment =""
        print(datetime.date.today())
        
        self.main_layout = QVBoxLayout()
        self.main_layout.setSpacing(5)
        self.main_layout.setContentsMargins(5,5,5,5)

        #date_line_layout = QHBoxLayout() date_widget = QDateEdit()
        top_lab = QLabel("Консультативное заключение врача травматолога-ортопеда")
        top_lab.setObjectName("heading")
        top_lab.setStyleSheet('''
                font-family: Courier New, monospace;
                text-align: center;
                font-size:16px;
                height: 50px;

                ''')
        top_lab.setMinimumHeight(50)
        self.main_layout.addWidget(top_lab, alignment=Qt.AlignHCenter)
        self.all_objects.append([top_lab])
        dict_to_building = []

        date_widget = QDateEdit()
        date_widget.setStyleSheet('''font-family: Courier New, monospace; font-size:16px;''')
        date_widget.setObjectName("date")
        date_widget.setDate(datetime.date.today())
        
        name = QLineEdit()
        name.setStyleSheet('''
                font-family: Courier New, monospace;
                font-size:16px;
                height: 50px;
                ''')
        name.setMaximumHeight(30)
        name.setObjectName("name")
        
        localis = QTextEdit()
        localis.setObjectName("textedit_loc")
        manipulation = Text_edit()
        manipulation.setObjectName("textedit_man")
        manipulation.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=manipulation.toPlainText()))
        diagnosis = Text_edit()
        diagnosis.setObjectName("textedit_diag")
        diagnosis.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=diagnosis.toPlainText()))
        appointment = Text_edit()
        appointment.setObjectName("textedit_appo")
        appointment.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=appointment.toPlainText()))
        recomendation = Text_edit()
        recomendation.setObjectName("textedit_recom")
        recomendation.mouseDoubleClick.connect(lambda: self.text_edit_double_click_process(text=recomendation.toPlainText()))
        

        for item in [["Дата", date_widget],["Ф. И. О.",name],["Локальный статус:", localis], ["Манипуляции:", manipulation], 
                    ["Диагноз:", diagnosis], ["Назначенные\nобследования:", appointment],[ "Рекомендации:", recomendation]]:
            lab = QLabel(item[0], parent=item[1])
            #lab.adjustSize()
            if item[0] in ["Дата", "Ф. И. О."]:
                lab.setObjectName("lab_" + item[1].objectName())
            else:
                lab.setObjectName("lab" + item[1].objectName()[item[1].objectName().find("_"):])
            lab.setStyleSheet('''
                font-family: Courier New, monospace;
                font-size:16px;
                height: 20px;
                ''')
            if "textedit" in item[1].objectName():
                textedit=item[1]
                if item[0] == "Локальный статус:":

                    self.localis = textedit
                elif item[0] == "Манипуляции:":
                    self.manipulation = textedit
                elif item[0] == "Диагноз:":
                    self.diagnosis = textedit
                elif item[0] == "Назначенные\nобследования:":
                    self.appointment = textedit
                    
                elif item[0] == "Рекомендации:":
                    self.recomendation = textedit
                #textedit.setMaximumHeight(60)
                textedit.setWordWrapMode(1)
                textedit.setStyleSheet('''
                    font-family: Courier New, monospace;
                    font-size:16px;

                    ''')
                #textedit.setSizeAdjustPolicy(0)
                textedit.setWordWrapMode(1)
                #textedit.textChanged.connect(self.textedit_resizing)


            #item[1].setMaximumWidth(self.width()-lab.width())
            #item[1].setMinimumWidth(self.width()-lab.width())
            #print("consult",self.width()-lab.width())
            #print("consult",self.width())
            #print("consult",lab.width())
            dict_to_building.append([lab, item[1]])

        inc = 0
        for i in dict_to_building:

            self.all_objects.append([i[0],i[1]])
            line = QHBoxLayout()
            
            line.addWidget(i[0])
            line.addWidget(i[1])
            if inc in [0]:
                line.addStretch(1)
            #line.addStretch(1)
            #widget.adjustSize()
            inc+=1



            self.main_layout.addLayout(line)

        foot_lab = QLabel("Врач травматолог-ортопед                                        (А.В.Шакулин)")
        foot_lab.setObjectName("footer")
        foot_lab.setStyleSheet('''font-family: Courier New, monospace; font-size:16px;''')
        self.main_layout.addWidget(foot_lab)
        self.main_layout.insertSpacing(-1, 30)
        self.all_objects.append([foot_lab])


        but_layout = QHBoxLayout()
        button = PushButton("Сохранить в файле .docx")
        button.clicked.connect(self.text_proceess)
        button.color = "#ffb4a2"
        button.font_size = 14
        
        button_blank = PushButton("Сохранить на бланке .docx")
        button_blank.setObjectName("save_in_blank")
        button_blank.clicked.connect(self.text_proceess)
        button_blank.color = "#ffb4a2"
        button_blank.font_size = 14

        button_clear = PushButton("Очистить форму")
        button_clear.clicked.connect(self.clear_buttom_push)
        button_clear.color = '#fa8090'
        button_clear.font_size = 14
        but_layout.addWidget(button)
        but_layout.addWidget(button_blank)
        but_layout.addWidget(button_clear)
        self.main_layout.addLayout(but_layout)
        self.main_layout.addStretch(1)

        self.setLayout(self.main_layout)
    
    def text_edit_double_click_process(self, text):
        obj = self.sender()
        print(obj.objectName())
        with open('samples_manipulation.pickle',  'rb') as f:
            manipulation  = pickle.load(f)

        with open('samples_diagnosis.pickle',  'rb') as f:
            diagnosis  = pickle.load(f)
        
        with open('samples_appointment.pickle', 'rb') as z:
            appointment= pickle.load(z)
        
        with open('samples_recomendation.pickle', 'rb') as x:
            recomendation= pickle.load(x)
        


        if "diag" in obj.objectName():
            self.sample = SampleWindow (source=diagnosis, source_name='samples_diagnosis.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
        elif "man" in obj.objectName():
            self.sample = SampleWindow (source=manipulation, source_name='samples_manipulation.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
        elif "appo" in obj.objectName():
            self.sample = SampleWindow (source=appointment, source_name='samples_appointment.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
        elif "recom" in obj.objectName():
            self.sample = SampleWindow (source=recomendation, source_name='samples_recomendation.pickle', text = text)
            self.sample.sample_window_text_sending.connect(lambda: self.text_pasting(obj, self.sample.left_box.toPlainText()))
            self.sample.show()
    
    def text_pasting(self, obj, text):
        if text !="":
            obj.setText(text)


    def textedit_resizing(self):
        textedit = self.sender()
        if textedit.isVisible():
            
            label = ""
            for line in self.all_objects:
                if textedit in line:
                    label = line[0]

            
            w_sender = textedit.width()
            h_sender = textedit.height()
            #print(self.sender())
            #print(textedit.objectName())
            #print("resizing")
            #print(str(self.sender().toPlainText()))
            other_wid_width = label.width()
            
            #print("other_wid_width", other_wid_width)

            font = self.font
            #textedit.document().defaultFont()
            fontMetrics = QFontMetrics(font)
            print(other_wid_width)
            textSize = fontMetrics.size(0, textedit.toPlainText())
            w = textSize.width() + 15
            h = textSize.height() +5
            if w > self.width() - other_wid_width:
                w = self.width() - other_wid_width-20
            else:
                pass
            textedit.setMinimumSize(w, h)
            textedit.setMaximumSize(w, h)
            textedit.resize(w, h)
            print(textedit)
            print(w, h)
            print(self.width())
        

    def text_proceess(self):
        desktop_destination = os.path.join(os.environ['USERPROFILE'], "Desktop\\ORTHO_STATUS\\")
        print(desktop_destination)
        try:
            os.mkdir(desktop_destination, mode=stat.S_ISUID)
        except OSError as error:
            print(error)   
        
        fileName_png = ""

        if self.sender().objectName() == "save_in_blank":
            options = QFileDialog.Options()
            options |= QFileDialog.DontUseNativeDialog
            fileName_png, _ = QFileDialog.getOpenFileName(self,"Выберите файл с изображением в формате .png для фона в заключении",directory=desktop_destination, 
                                                        filter="Text Files (*.png)", options=options)
            
            #f = open(fileName, 'rb')



        #mydoc.add_paragraph("Дата %s, время: %s" %(time.strftime("%d-%m-%Y"), time.strftime("%H:%M")))
        
        #"textedit_loc"
        #"textedit_man"
        #"textedit_diag"
        #"textedit_appo"
        #"textedit_recom"
        '''
        heading
        lab_date
        date
        lab_name
        name
        lab_loc
        textedit_loc
        lab_man
        textedit_man
        lab_diag
        textedit_diag
        lab_appo
        textedit_appo
        lab_recom
        textedit_recom
        footer
        '''
        doc_prepare = docx.Document()

        for line in self.all_objects:
            
            par = doc_prepare.add_paragraph(style="No Spacing")
            paragraph_format = par.paragraph_format
            paragraph_format.left_indent = Cm(1.5)
            
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


            for item in line:
                
                par = doc_prepare.add_paragraph(style="No Spacing")
                paragraph_format = par.paragraph_format
                #paragraph_format.left_indent = Cm(1.5)
                
                if item.objectName() in ["heading", "footer"]:
                    paragraph_format = par.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    text = item.text()[:]
                    run = par.add_run(text)
                    run.bold = True
                    run.underline = True
                    font = run.font
                    #font.name = 'Bad Script'
                    font.name = 'Microsoft JhengHei Light'
                    
                    
                    if item.objectName() == "footer":
                        prev_par = doc_prepare.paragraphs[-2]
                        prev_par.add_run("\n")
                        font.size = Pt(10)
                    else:
                        paragraph_format.space_before = Cm(3.0)
                        font.size = Pt(12)
                    continue
                
                
                


                print(item.objectName())
                if "lab" in item.objectName() or item.objectName() in ["date", "name"]:

                    text = item.text()[:]
                    if "\n" in text:
                        text = text[:text.find("\n")] + " " + text[text.find("\n")+1:]
                    print(text)
                    if item.objectName() in ["date", "name"]:
                        par = doc_prepare.paragraphs[-2]
                        run = par.add_run("  " + text)
                        font = run.font
                        #font.name = 'Bad Script'
                        font.name = 'Microsoft JhengHei Light'
                        font.size = Pt(10)
                    else:
                        run = par.add_run(text)
                        run.bold = True
                        run.underline = True
                        font = run.font
                        #font.name = 'Bad Script'
                        font.name = 'Microsoft JhengHei Light'
                        font.size = Pt(10)
                    #print(text)
                

                elif "textedit" in item.objectName():
                   
                    text_1 = item.toPlainText()
                    print("text_1",text_1)
                    if text_1 not in ["", "нет"]:
                        run = par.add_run(text_1)
                        font = run.font
                        #font.name = 'Bad Script'
                        font.name = 'Microsoft JhengHei Light'
                        if item.objectName() in ["textedit_loc", "textedit_diag", "textedit_man"]:
                            font.size = Pt(9)
                        else:
                            font.size = Pt(12)
                        #print(item.toPlainText())
                    else:
                        for paragraph in [doc_prepare.paragraphs[-2], doc_prepare.paragraphs[-1]]:
                            p = paragraph._element
                            p.getparent().remove(p)
                            paragraph._p = paragraph._element = None

                        #mydoc.paragraphs[-2].text = None
                        #mydoc.paragraphs[-1].text = None
                        
                        #par.runs[-2].text = ""
                        
                #par.add_run(" ")
            for paragraph in doc_prepare.paragraphs:
                if paragraph.text == "":
                    p = paragraph._element
                    p.getparent().remove(p)
                    paragraph._p = paragraph._element = None
        
        if fileName_png !="":
            p = doc_prepare.paragraphs[0]

            add_float_picture(p, fileName_png, width=Cm(18.0), pos_x=Pt(20), pos_y=Pt(30))

        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self,"Сохранение файла .docx",directory=desktop_destination, 
                                                    filter="All Files (*);;Text Files (*.docx)", options=options)
        if fileName:
            print(fileName)
        if fileName[-5:] !=".docx":
            full_way = fileName +".docx"
        else:
            full_way = fileName
        print(full_way)
        
        doc_prepare.save(full_way)

    
    def setLocalis(self, localis="", manipulation="", diagnosis="", appointment="", recomendation=""):
        if localis !="": 
            print(localis) 
            self.localis.setText(localis)
        if manipulation !="": self.manipulation.setText(manipulation)
        if diagnosis !="": self.diagnosis.setText(diagnosis)
        if appointment !="": self.appointment.setText(appointment)
        if recomendation !="": self.recomendation.setText(recomendation)

    def clear_buttom_push(self):
        self.clear_form_signal.emit()
        self.clear_form()

    def clear_form(self):

        layout = self.layout()
        print(layout)
        self.clearLayout(layout)
        import sip
        sip.delete(layout)
        self.initUI()
        #for line in self.all_objects:
        #    for item in line:
        #        if "textedit" in item.objectName():
        #            item.setText()

    def clearLayout(self, layout):
        if layout != None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget() is not None:
                    child.widget().deleteLater()
                elif child.layout() is not None:
                    self.clearLayout(child.layout())


class Label_press(QLabel):
    label_pressed = pyqtSignal()
    label_double_pressed = pyqtSignal()
    def __init__(self):
        super().__init__()
    
    def mousePressEvent(self,event):
        self.label_pressed.emit()
        print("press")

        return QLabel.mousePressEvent(self,event)



class SampleWindow(QWidget):
    sample_window_text_sending = pyqtSignal()
    def __init__(self,source, source_name, text, parent=None):
        super().__init__()
        self.font = QFont("Courier New", pointSize=14)
        self.source = source
        self.source_name = source_name
        self.text = text
        self.initUI()

    def initUI(self):
        
        self.add_group_box = ""

        main_layout = QGridLayout(self)
        main_layout.addWidget(QLabel("Текстовое поле"), 0,0)
        main_layout.addWidget(QLabel("Шаблоны"), 0, 1)

        self.left_box = QTextEdit()
        if self.text!="":
            self.left_box.setText(self.text + "\n")
        main_layout.addWidget(self.left_box, 1,0,1,1)
        self.right_box = QVBoxLayout()
        self.widget = QWidget()
        self.widget.setLayout(self.right_box)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(self.widget)
        main_layout.addWidget(scroll,1,1,1,1)
        #but_edit = PushButton("Редактировать список")
        #but_edit.clicked.connect(self.table_display)
        #but_edit.color = "#ffb4a2"
        #but_edit.font_size = 14

        but_ok = PushButton("Ok")
        but_ok.clicked.connect(self.close_sample_window)
        but_ok.color = "#adff2f"
        but_ok.font_size = 14

        #main_layout.addWidget(but_edit,2,1,1,1)
        main_layout.addWidget(but_ok,2,0,1,2)
        buttons= QHBoxLayout()
        self.resize(640,480)

        width = self.widget.width()
        print(self.width())
        
        self.labels_samples = []
        self.but_ok_list = []
        #print("width", width)
        
        self.building_element_tree()

        self.add_editing_chapter_key()
        




        self.show()
    

    def building_element_tree(self):
        self.clearLayout(self.right_box)
        self.labels_samples= []
        self.but_ok_list=[]
        
        for key, value in self.source.items():
            group_box = QGroupBox(key)
            group_box.setStyleSheet('''
                QGroupBox {
                    margin-top: 2ex;
                    font-size: 14px;
                    
                }
                QGroupBox:enabled {
                    border: 3px solid #CA8372;
                    border-radius: 5px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 3ex;
                }
            ''')
            if isinstance(value, list):
                but_layout = QVBoxLayout()
                but_layout.setSpacing(2)
                increment = 0
                for item in value:
                    line_box = QHBoxLayout()
                    press_lab = Label_press()
                    if len(item) > 40:
                        text_new = ""
                        inc = 1
                        for i in item:
                            text_new+=i
                            inc+=1
                            if inc ==41:
                                text_new+="\n"
                                inc = 1
                        press_lab.setText(text_new)
                    else:
                        press_lab.setText(item)
                    self.labels_samples.append(press_lab)
                    press_lab.setObjectName("lab_" + key +"}"+ str(increment))
                    press_lab.setStyleSheet('''border: 1px solid #CA8372; padding: 4 px; border-radius: 3px''')
                    press_lab.label_pressed.connect(self.label_pressed_procees)
                    #button.setMaximumWidth(100)
                    edit_layout = QVBoxLayout()
                    but_edit = QPushButton("Edit")
                    but_edit_ok = QPushButton("Ok")
                    self.but_ok_list.append(but_edit_ok)
                    but_edit.clicked.connect(self.label_editing)
                    #but_edit_ok.clicked.connect(self.but_ok_proccess)
                    but_edit.setObjectName("but_" + key +"}"+ str(increment))
                    but_edit_ok.setObjectName("but_ok_" + key  +"}" + str(increment))
                    line_box.addWidget(press_lab)
                    edit_layout.addWidget(but_edit)
                    edit_layout.addWidget(but_edit_ok)
                    line_box.addLayout(edit_layout)
                    but_layout.addLayout(line_box)
                    

                    but_edit_ok.hide()
                    print(press_lab.width())
                    increment+=1
                but_add = QPushButton("Добавить элемент")
                but_add.setObjectName("but_add_" + key)
                but_add.clicked.connect(self.adding_press_lab)
                but_layout.addWidget(but_add)
            
            group_box.setLayout(but_layout)
            self.right_box.addWidget(group_box)




    def adding_press_lab(self):
        text, ok = QInputDialog.getText(self, 'Input Dialog', 'Enter text:')
        if ok:
            self.source[self.sender().objectName()[8:]].append(text)
            with open(self.source_name, 'wb') as f:
                    pickle.dump(self.source, f)
            self.building_element_tree()
            self.add_editing_chapter_key()



    def add_editing_chapter_key(self):
        if self.add_group_box != "":
            self.add_group_box.setParent(None)
            self.add_group_box = ""
        
        self.add_group_box = QGroupBox("Редактировать раздел")
        self.add_group_box.setStyleSheet('''
                QGroupBox {
                    margin-top: 2ex;
                    font-size: 14px;
                    
                }
                QGroupBox:enabled {
                    border: 3px solid #CA8372;
                    border-radius: 5px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 3ex;
                }
            ''')
        self.button_add = QPushButton("Довавить")
        self.button_add.setObjectName("add")
        self.button_del_chap = QPushButton("Удалить раздел")
        self.button_del_chap.setObjectName("del")
        self.button_del_chap.clicked.connect(self.chapter_edit_proccess)
        self.button_add.clicked.connect(self.chapter_edit_proccess)
        

        but_lay_chapter = QHBoxLayout()
        but_lay_chapter.addWidget(self.button_add)
        but_lay_chapter.addWidget(self.button_del_chap)
        self.add_group_box.setLayout(but_lay_chapter)

        self.right_box.addWidget(self.add_group_box)




    def label_editing(self):
        but_name = self.sender().objectName()
        label = ""
        text = ""
        key = but_name[4:but_name.find("}")]
        label_index = int(but_name[but_name.find("}")+1:])
        print(label_index)
        for lab in self.labels_samples:
            print(lab.objectName())
            if lab.objectName()[4:] == but_name[4:]:
                print(lab.objectName())
                label = lab
                text = lab.text()
        print(text)
        print(key)
        text_input, ok = QInputDialog.getMultiLineText(self, 'Редактирование элемента', 'Откорректируйте текст', text)
        if ok and text_input!="":
            text_input = text_input.replace("\n", "")
            
            self.source[key][label_index] = text_input 
            with open(self.source_name, 'wb') as f:
                pickle.dump(self.source, f)
            self.building_element_tree()
            self.add_editing_chapter_key()

        elif ok and text_input =="":
            self.source[key].pop(label_index)
            with open(self.source_name, 'wb') as f:
                pickle.dump(self.source, f)
            self.building_element_tree()
            self.add_editing_chapter_key()


    

    def chapter_edit_proccess(self):
        if self.sender().objectName() == "add":
            text, ok = QInputDialog.getText(self, 'Input Dialog', 'Enter text:')
            if ok:
                self.source[text] = []
                with open(self.source_name, 'wb') as f:
                    pickle.dump(self.source, f)
                self.building_element_tree()
                self.add_editing_chapter_key()

        elif self.sender().objectName() == "del":
            if self.source.keys() != []:
                item, ok = QInputDialog.getItem(self, 'Input Dialog', 'Enter text:', self.source.keys(), 0)
                if ok:
                    self.source.pop(item)
                    with open(self.source_name, 'wb') as f:
                        pickle.dump(self.source, f)
            else:

                msg = QMessageBox()
                msg.setWindowTitle("Информация")
                msg.setText("Информация не введена!!!")
                msg.setIcon(QMessageBox.Information)
                msg.show()
                msg.exec_()


    def editing_proccess(self, lab, texted):
        print(texted.toPlainText())
        lab.setText(texted.toPlainText())
        texted.setParent(None)
        self.sender().hide()



    def table_display(self):
        self.table = Table_view(self.source)
        self.table.show()

    def close_sample_window(self):
        self.sample_window_text_sending.emit()
        self.close()

    def label_pressed_procees(self):
        self.sender().text()
        text_cleaned = self.sender().text().replace("\n","")
        if self.left_box.toPlainText() != "":
            self.left_box.setText(self.left_box.toPlainText() + "\n" + text_cleaned)
        else:
            self.left_box.setText(self.left_box.toPlainText() +text_cleaned)

    def clearLayout(self, layout):
        if layout != None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget() is not None:
                    child.widget().deleteLater()
                elif child.layout() is not None:
                    self.clearLayout(child.layout())